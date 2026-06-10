import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.utils import timezone

def set_cell_background(cell, hex_color):
    """
    Utility to color-fill a table cell background.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def style_row_height(row, height_pt):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_pt * 20))) # Pt to twips conversion
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

def generate_docx_report(assessment, report_type):
    doc = Document()
    
    # Configure default style font to Segoe UI/Calibri
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"RISKPILOT CYBER RISK REPORT")
    title_run.font.size = Pt(9)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x88, 0x22, 0x11) # Rust Red

    # Heading 1
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run()
    h1_run.font.name = 'Segoe UI'
    h1_run.font.size = Pt(20)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(0x00, 0x33, 0x88) # Deep Blue

    if report_type == 'ExecutiveSummary':
        h1_run.text = f"Executive Summary Assessment: {assessment.name}"
        
        # Meta info block
        p_meta = doc.add_paragraph()
        p_meta.add_run("Client Organisation: ").bold = True
        p_meta.add_run(f"{assessment.client.name}\n")
        p_meta.add_run("Client Contact: ").bold = True
        p_meta.add_run(f"{assessment.client.email}\n")
        p_meta.add_run("Generated Date: ").bold = True
        p_meta.add_run(f"{timezone.now().strftime('%Y-%m-%d %H:%M')}\n")
        p_meta.add_run("Framework / Methodology: ").bold = True
        p_meta.add_run(f"{assessment.methodology_version}\n")
        p_meta.add_run("Current Assessment Status: ").bold = True
        p_meta.add_run(f"{assessment.status}")

        doc.add_heading("Assessment Overview", level=2)
        doc.add_paragraph(
            f"This executive summary outlines the key risk findings, assets assessed, and overall risk posture discovered during "
            f"the evaluation of '{assessment.name}'. The assessment is conducted under the {assessment.methodology_version} rules."
        )

        doc.add_heading("Risk Posture Summary", level=2)
        total_risks = assessment.risk_items.count()
        high_risks = len([r for r in assessment.risk_items.all() if r.risk_category in ['High', 'Critical']])
        open_findings = assessment.findings.filter(status='Open').count()
        mitigated_risks = assessment.risk_items.filter(treatment__status__in=['Mitigated', 'Closed']).count()

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Risk Metric"
        hdr_cells[1].text = "Count / Value"
        for cell in hdr_cells:
            set_cell_background(cell, "003388")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell.paragraphs[0].runs[0].font.bold = True

        metrics = [
            ("Total Identified Risks", str(total_risks)),
            ("High & Critical Severity Risks", str(high_risks)),
            ("Open Vulnerability Gaps / Findings", str(open_findings)),
            ("Successfully Mitigated Risks", str(mitigated_risks)),
        ]

        for m_name, m_val in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = m_name
            row_cells[1].text = m_val
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    elif report_type == 'DetailedRiskAssessment':
        h1_run.text = f"Detailed Risk Assessment: {assessment.name}"
        
        # Details
        doc.add_heading("Client and Scope", level=2)
        p_scope = doc.add_paragraph()
        p_scope.add_run("Client: ").bold = True
        p_scope.add_run(f"{assessment.client.name} ({assessment.client.email})\n")
        p_scope.add_run("Methodology: ").bold = True
        p_scope.add_run(f"{assessment.methodology_version}\n")
        p_scope.add_run("Main System / Asset: ").bold = True
        p_scope.add_run(f"{assessment.asset} (Location: {assessment.location}, Owner: {assessment.owner})")

        doc.add_heading("Risk Items Register Details", level=2)
        for idx, item in enumerate(assessment.risk_items.all(), start=1):
            doc.add_heading(f"Risk #{idx}: {item.asset_name}", level=3)
            p_desc = doc.add_paragraph()
            p_desc.add_run("Threat Scenario: ").bold = True
            p_desc.add_run(f"{item.threat.name} - {item.threat.description}\n")
            p_desc.add_run("Vulnerability: ").bold = True
            p_desc.add_run(f"{item.vulnerability}\n")
            p_desc.add_run("Existing Controls: ").bold = True
            p_desc.add_run(f"{item.existing_controls}\n")
            
            p_scores = doc.add_paragraph()
            p_scores.add_run("Inherent Risk Score: ").bold = True
            p_scores.add_run(f"{item.risk_score} ({item.risk_category})\n")
            p_scores.add_run("Residual Risk Score: ").bold = True
            p_scores.add_run(f"{item.residual_risk_score if item.residual_risk_score is not None else 'N/A'} ({item.residual_risk_category})")

            # Link findings if any
            findings = item.findings.all()
            if findings.exists():
                doc.add_paragraph("Linked Findings / Vulnerability Gaps:").runs[0].font.bold = True
                for f in findings:
                    doc.add_paragraph(f"• {f.title} - Severity: {f.severity} (Status: {f.status})", style='List Bullet')
            doc.add_paragraph()

    elif report_type == 'RiskRegister':
        h1_run.text = f"Risk Register Report: {assessment.name}"
        doc.add_paragraph("This report contains the raw tabular cyber risk registry for compliance auditing and governance reviews.")

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["Asset", "Threat Scenario", "Inherent Score", "Inherent Category", "Residual Score", "Residual Category"]
        for c_idx, h in enumerate(headers):
            hdr_cells[c_idx].text = h
            set_cell_background(hdr_cells[c_idx], "003388")
            hdr_cells[c_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            hdr_cells[c_idx].paragraphs[0].runs[0].font.bold = True

        for item in assessment.risk_items.all():
            row_cells = table.add_row().cells
            row_cells[0].text = item.asset_name
            row_cells[1].text = item.threat.name
            row_cells[2].text = str(item.risk_score)
            row_cells[3].text = item.risk_category
            row_cells[4].text = str(item.residual_risk_score if item.residual_risk_score is not None else "N/A")
            row_cells[5].text = item.residual_risk_category

    elif report_type == 'TreatmentPlan':
        h1_run.text = f"Risk Treatment Plan: {assessment.name}"
        doc.add_paragraph("Remediation and risk mitigation treatments assigned to identified security vulnerabilities.")

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["Asset Affected", "Mitigation Action Required", "Owner", "Target Date", "Status"]
        for c_idx, h in enumerate(headers):
            hdr_cells[c_idx].text = h
            set_cell_background(hdr_cells[c_idx], "003388")
            hdr_cells[c_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            hdr_cells[c_idx].paragraphs[0].runs[0].font.bold = True

        for item in assessment.risk_items.all():
            treatment = getattr(item, 'treatment', None)
            if treatment:
                row_cells = table.add_row().cells
                row_cells[0].text = item.asset_name
                row_cells[1].text = treatment.action
                row_cells[2].text = treatment.owner
                row_cells[3].text = treatment.target_date.strftime("%Y-%m-%d") if treatment.target_date else "N/A"
                row_cells[4].text = treatment.status

    elif report_type == 'ChangeRequest':
        h1_run.text = f"Change Request Assessment: {assessment.name}"
        
        doc.add_heading("Change Context", level=2)
        doc.add_paragraph(assessment.change_request or "No associated change request details provided.")

        doc.add_heading("Business Process Impact", level=2)
        doc.add_paragraph(assessment.business_process_impact or "No impact assessment provided.")

        doc.add_heading("CIA Security Impact Assessment", level=2)
        doc.add_paragraph(f"• Confidentiality Affected: {'Yes' if assessment.confidentiality_affected else 'No'}", style='List Bullet')
        doc.add_paragraph(f"• Integrity Affected: {'Yes' if assessment.integrity_affected else 'No'}", style='List Bullet')
        doc.add_paragraph(f"• Availability Affected: {'Yes' if assessment.availability_affected else 'No'}", style='List Bullet')

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
